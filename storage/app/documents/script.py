from pathlib import Path
from docxtpl import DocxTemplate, InlineImage
from docx import Document
import base64
import json
import re
import os
import argparse
import io
import cairosvg

DOCX_FILTER="docx:'Office Open XML Text'"
HTML_FILTER="HTML"
PDF_FILTER="pdf:writer_pdf_Export"

def remove_headers_and_placeholder_symbols(fileName: str)->None:

    fileNoHeader = str(Path(fileName).stem) + '-no-header.docx'
    document = Document(fileName)
    for section in document.sections:
        section.header.is_linked_to_previous = True
        section.footer.is_linked_to_previous = True

    for par in document.paragraphs:
        par.text = re.sub(" ?([... …․﹒．̲ _]{3,}\s*?)\n*", ' {{ placeholder }} ', par.text)

    for idx, table in enumerate(document.tables):
        for idx2, row in enumerate(table.rows):
            for idx3, cell in enumerate(row.cells):
                for para in (cell.paragraphs):
                    if not para.text or para.text=='\n':
                        para.text = ' {{ placeholder }} '
    document.save(fileNoHeader)

def set_template_fields(fileNameOriginal: str, fileNameFromHtml: str)->None:
    document = Document(fileNameOriginal)
    document2 = Document(fileNameFromHtml)

    if len(document.paragraphs) < len(document2.paragraphs):
        for x in range(len(document2.paragraphs) - len(document.paragraphs)):
            document.add_paragraph()

    for (para1, para2) in zip(document.paragraphs, document2.paragraphs):
        if  not para1.style:
            para1.style = para2.style
        para1.text = para2.text

    for idx, table in enumerate(document.tables):
        for idx2, row in enumerate(table.rows):
            for idx3, cell in enumerate(row.cells):
                if cell.text != document2.tables[idx].rows[idx2].cells[idx3]:
                    cell.text = document2.tables[idx].rows[idx2].cells[idx3].text
    document.save(fileNameOriginal)


def convert_to(file: str, filter: str)->None:
    flt = ''
    if (filter == 'doc'):
        flt = DOCX_FILTER
    elif (filter =='html'):
        flt = HTML_FILTER
    elif (filter=='pdf'):
        flt = PDF_FILTER
    process = os.system(f"libreoffice --headless --convert-to {flt} {file}")


def fill_in_document(fileName: str, data: str)->None:
    document = DocxTemplate(fileName)
    fromFile = ''
    with open(data, "r") as f:
        fromFile = json.loads(f.read())
    context = {}
    files = []
    for key in fromFile.keys():
        if "data" in fromFile[key]:
            svg_bytes = base64.b64decode(fromFile[key][26:])
            svg_io = io.BytesIO(svg_bytes)
            png_out = cairosvg.svg2png(bytestring=svg_io.read())
            with open(key+'.png', "wb") as fh:
                fh.write(png_out)
                fh.close()
            # with open(key+'.svg', "wb") as fh:
            #     # fh.write(fromFile[key][26:])
            #     fh.write(base64.b64decode(fromFile[key][26:]))
            #     fh.close()
            files.append(key+'.png')
            context[key] = InlineImage(document, key+'.png')
            # fromFile[key] = InlineImage(document, fromFile[key][26:])
            # val = val[26:]
            # print(fromFile[key])
        else:
            context[key] = fromFile[key]
    document.render(context)
    document.save('FILLED_IN-'+fileName)
    for f in files:
        os.unlink(f)


def main():
    parser = argparse.ArgumentParser(description="Document Processing Tool")
    subparsers = parser.add_subparsers(dest="subparser_name")

    parser_remove = subparsers.add_parser("remove-header-symbol", help="Remove headers and placeholder symbols")
    parser_remove.add_argument("file", type=str, help="Input file name")

    parser_set_template = subparsers.add_parser("set-fields", help="Set template fields")
    parser_set_template.add_argument("original", type=str, help="Original file name")
    parser_set_template.add_argument("html", type=str, help="HTML-based template file name")

    parser_convert = subparsers.add_parser("convert-to", help="Convert a file to different formats")
    parser_convert.add_argument("file", type=str, help="Input file name")
    parser_convert.add_argument("format", type=str, choices=['html', 'doc', 'pdf'], help="Output format")

    parser_fill_in = subparsers.add_parser("fill-in", help="Fill in a document with data")
    parser_fill_in.add_argument("file", type=str, help="Input file name")
    parser_fill_in.add_argument("data", type=str, help="Data file name")

    args = parser.parse_args()

    if args.subparser_name == "remove-header-symbol":
        remove_headers_and_placeholder_symbols(args.file)
    elif args.subparser_name == "set-fields":
        set_template_fields(args.original, args.html)
    elif args.subparser_name == "convert-to":
        convert_to(args.file, args.format)
    elif args.subparser_name == "fill-in":
        fill_in_document(args.file, args.data)

if __name__ == "__main__":
    main()
